import os, random
from datetime import datetime, timedelta
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "test_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def rd(sy=2024, ey=2026):
    s = datetime(sy,1,1); e = datetime(ey,3,1)
    return s + timedelta(days=random.randint(0,(e-s).days))

def doc(): return Document()
def h(d,t,l=0): return d.add_heading(t,level=l)
def p(d,t): return d.add_paragraph(t)
def tbl(d,hds,rows):
    t=d.add_table(rows=1+len(rows),cols=len(hds)); t.style='Table Grid'
    for i,x in enumerate(hds): t.rows[0].cells[i].text=x
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row): t.rows[ri+1].cells[ci].text=str(v)
    return t
def sv(d,fn): d.save(os.path.join(OUTPUT_DIR,fn))
